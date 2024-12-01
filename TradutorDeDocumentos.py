import requests
from docx import Document
import os

subscription_key = "YOUR_SUBSCRIPTION_KEY"
endpoint = "YOUR_ENDPOINT"
location = "eastus2"
language_destination = "pt-br"

def translate_document(path):
  document = Document(path) 
  full_text = []

  for para in document.paragraphs:
      translated_text = translate_text(para.text, language_destination)
      full_text.append(translated_text)

  translated_doc = Document()
  for line in full_text:
    print(line)
    translated_doc.add_paragraph(line)
  path_translated = path.replace('.docx', f'_{language_destination}.docx')
  translated_doc.save(path_translated)
  return path_translated
  
imput_file = '/content/Musica.docx'
translate_document(imput_file)


